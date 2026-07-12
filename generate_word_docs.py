"""
Step 4: Group reels by category and generate one Word document per category —
an overview of the common main points across all reels in that category,
plus a per-reel breakdown. Uses the same free Gemini API as step 3.

Usage:
    python generate_word_docs.py
"""
import json
import os
import re
import time
from collections import defaultdict
from pathlib import Path

import requests
from docx import Document

MANIFEST_PATH = Path("output/manifest.json")
OUTPUT_DIR = Path("output")
MODEL = "gemini-3.1-flash-lite"
API_URL = f"https://generativelanguage.googleapis.com/v1beta/models/{MODEL}:generateContent"


def synthesize_category_overview(api_key: str, category: str, entries: list[dict]) -> str:
    """Ask Gemini to write a short synthesized overview of the recurring
    main points across all reels in this category."""
    all_points = []
    for e in entries:
        all_points.extend(e.get("main_points", []))

    if not all_points:
        return "No detailed points were extracted for this category."

    points_text = "\n".join(f"- {p}" for p in all_points)
    prompt = (
        f"Here are main points collected from {len(entries)} Instagram reels, "
        f"all in the category '{category}':\n\n{points_text}\n\n"
        "Write a short synthesized overview (3-6 sentences) of the recurring "
        "themes and most important takeaways across these reels. Group related "
        "points together rather than just restating each one. Plain prose, no headers, "
        "no markdown formatting."
    )
    for attempt in range(5):
        resp = requests.post(
            f"{API_URL}?key={api_key}",
            json={"contents": [{"parts": [{"text": prompt}]}]},
            timeout=60,
        )
        if resp.status_code != 429:
            break
        wait = 2 ** attempt * 5
        print(f"    Rate limited, retrying in {wait}s (attempt {attempt+1}/5)...")
        time.sleep(wait)
    else:
        print("    Still rate limited after 5 retries — skipping overview.")
        return "Overview unavailable due to API rate limits."
    resp.raise_for_status()
    return resp.json()["candidates"][0]["content"]["parts"][0]["text"].strip()


def consolidate_categories(api_key: str, raw_categories: list[str]) -> dict[str, str]:
    """Ask Gemini to merge near-duplicate/overly-narrow category names (e.g.
    'AI Development', 'AI Productivity Tools', 'AI Productivity') into a
    smaller set of sensible canonical categories. Returns a mapping from
    every original category name to its canonical replacement."""
    if len(raw_categories) <= 1:
        return {c: c for c in raw_categories}

    categories_text = "\n".join(f"- {c}" for c in raw_categories)
    prompt = (
        "Here is a list of category names that were generated separately for "
        "different pieces of content, so there's a lot of near-duplication and "
        "overly narrow variants of what's really the same broader topic (e.g. "
        "'AI Development', 'AI Education', 'AI Productivity Tools', and "
        "'AI Productivity' should probably all become one category, like 'AI & Tech').\n\n"
        f"{categories_text}\n\n"
        "Group these into a smaller set of broader, sensible categories - merge "
        "obvious near-duplicates and overly specific variants, but don't over-merge "
        "genuinely distinct topics into one bucket (e.g. 'Cooking' and 'Career Advice' "
        "should stay separate). Respond with ONLY a JSON object mapping EVERY original "
        "category name (exactly as given, as the key) to its new canonical category name "
        "(as the value). Every input category must appear as a key exactly once."
    )
    for attempt in range(5):
        resp = requests.post(
            f"{API_URL}?key={api_key}",
            json={
                "contents": [{"parts": [{"text": prompt}]}],
                "generationConfig": {"responseMimeType": "application/json"},
            },
            timeout=60,
        )
        if resp.status_code != 429:
            break
        wait = 2 ** attempt * 5
        print(f"    Rate limited, retrying in {wait}s (attempt {attempt+1}/5)...")
        time.sleep(wait)
    else:
        print("    Still rate limited — using raw categories.")
        return {c: c for c in raw_categories}
    resp.raise_for_status()
    raw_text = resp.json()["candidates"][0]["content"]["parts"][0]["text"]
    mapping = json.loads(raw_text)

    # safety net: make sure every original category has a mapping, even if the
    # model missed one - falls back to keeping it as-is rather than dropping reels
    for c in raw_categories:
        mapping.setdefault(c, c)
    return mapping


def safe_filename(name: str) -> str:
    return re.sub(r"[^\w\-. ]", "_", name).strip() or "Uncategorized"


def build_doc(category: str, overview: str, entries: list[dict]) -> Document:
    doc = Document()
    doc.add_heading(category, level=0)

    doc.add_heading("Overview", level=1)
    doc.add_paragraph(overview)

    doc.add_heading(f"Reel-by-reel breakdown ({len(entries)} reels)", level=1)
    for e in entries:
        doc.add_heading(e.get("url", "Reel"), level=2)
        p = doc.add_paragraph()
        run = p.add_run(e.get("summary", ""))
        run.italic = True
        for point in e.get("main_points", []):
            doc.add_paragraph(point, style="List Bullet")

    return doc


def main():
    api_key = os.environ.get("GEMINI_API_KEY")
    if not api_key:
        raise SystemExit(
            "Set GEMINI_API_KEY first. Get a free key at https://aistudio.google.com/apikey\n"
            "  export GEMINI_API_KEY=your-key-here"
        )

    manifest = json.loads(MANIFEST_PATH.read_text(encoding="utf-8"))

    categorized_entries = [e for e in manifest.values() if e.get("categorized")]
    raw_categories = sorted({e.get("category", "Uncategorized") for e in categorized_entries})
    print(f"Found {len(raw_categories)} raw category names. Consolidating similar ones...")
    category_map = consolidate_categories(api_key, raw_categories)
    merged_count = len(raw_categories) - len(set(category_map.values()))
    print(f"Merged down to {len(set(category_map.values()))} categories "
          f"({merged_count} near-duplicates folded in).")

    by_category = defaultdict(list)
    for entry in categorized_entries:
        raw_category = entry.get("category", "Uncategorized")
        canonical_category = category_map.get(raw_category, raw_category)
        by_category[canonical_category].append(entry)

    print(f"Generating docs for {len(by_category)} categories...")
    OUTPUT_DIR.mkdir(exist_ok=True)

    for category, entries in by_category.items():
        print(f"  {category} ({len(entries)} reels)")
        overview = synthesize_category_overview(api_key, category, entries)
        doc = build_doc(category, overview, entries)
        out_path = OUTPUT_DIR / f"{safe_filename(category)}.docx"
        doc.save(out_path)
        time.sleep(4)

    print(f"Done. Word docs saved in {OUTPUT_DIR}/")


if __name__ == "__main__":
    main()
